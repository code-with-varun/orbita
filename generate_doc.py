import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = docx.Document()

# Set Margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Styles
title_style = doc.styles['Title']
title_style.font.name = 'Arial'
title_style.font.size = Pt(24)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(11, 24, 43)

h1_style = doc.styles['Heading 1']
h1_style.font.name = 'Arial'
h1_style.font.size = Pt(16)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(59, 130, 246)

p = doc.add_paragraph('Orbita Work Management System', style='Title')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

p_sub = doc.add_paragraph('4 Simplified Core Types Architecture & QA Verification Report')
p_sub.runs[0].font.size = Pt(14)
p_sub.runs[0].font.color.rgb = RGBColor(100, 116, 139)

doc.add_heading('1. Executive Architecture Summary', level=1)
doc.add_paragraph(
    'In strict alignment with the Simplifications sheet in the Orbita Project BRD, the entire system '
    'has been streamlined into 4 core types: Task, Routine, Goal, and Project. Each type operates with '
    'Personal and Work workspaces, automatic Eisenhower Priority Quadrant assignment (Q1 to Q4), '
    'and dedicated life cycle execution rules.'
)

doc.add_heading('2. The 4 Simplified Orbita Types', level=1)

table_m = doc.add_table(rows=1, cols=6)
table_m.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_m = table_m.rows[0].cells
headers_m = ['Orbita Type', 'Time Tracking', 'Workspace', 'Decided By', 'Life Cycle', 'Representative Examples']
for i, h in enumerate(headers_m):
    hdr_m[i].text = h
    hdr_m[i].paragraphs[0].runs[0].font.bold = True

matrix_items = [
    ('TASK', 'No', 'Personal / Work', 'One-time completion single occurrence, simple completion', 'Create -> Complete', 'Buy groceries, post a photo, search document'),
    ('ROUTINE', 'No', 'Personal / Work', 'Automatic recurrence, repeats automatically', 'Repeat -> Complete Occurrence', 'Weekly meeting, EB bill, nth day of month/week'),
    ('GOAL', 'Yes', 'Personal / Work', 'Time-tracked ongoing effort, multiple focus timer sessions, flat structure', 'Create -> Start Timer -> Stop/Pause -> Complete', 'Learning JavaScript, YouTube content creation, cooking, personal accounting'),
    ('PROJECT', 'Yes', 'Personal / Work', 'Structured multistep delivery, needs stage -> task hierarchy', 'Stage -> Task -> Complete Project (all tasks must finish)', 'Building CRM app, MERN project, Website development')
]

for row in matrix_items:
    r_cells = table_m.add_row().cells
    for c_i, val in enumerate(row):
        r_cells[c_i].text = val

doc.add_heading('3. System Verification & QA Checklist', level=1)

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
headers = ['Component / Module', 'Feature Under Test', 'Status', 'Verification Note']
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True

items = [
    ('4 Core Orbita Types', 'Tasks, Routines, Goals, Projects Architecture', 'PASSED', 'Recreated schema and UI for the 4 pure types'),
    ('Workspace Division', 'Personal vs Work 2-Way Filtering', 'PASSED', 'Seamless toggle between Personal, Work, and All items'),
    ('Project Hierarchy', 'Stage -> Task Multistep Delivery', 'PASSED', 'Stage-tasks progress bars; project auto-completes only when all tasks are done'),
    ('Goal Focus Tracking', 'Live Stopwatch Timer & Hours Target', 'PASSED', 'Live real-time ticking stopwatch badge, effort target progress, and timesheets'),
    ('Routine Recurrence', 'Automated Recurrence Patterns', 'PASSED', 'Daily, Weekly, Monthly, and Yearly recurrence schedules'),
    ('Clean State Reset', 'Wipe Database to 0 Users / 0 Items', 'PASSED', 'Pure blank slate ready for fresh user registration and login'),
    ('Eisenhower Matrix', '2x2 Decision Grid (Q1 to Q4)', 'PASSED', 'Automatic assignment into Q1: Do Now, Q2: Plan, Q3: Quick Action, Q4: Optional'),
    ('Monthly Highlights', 'Scorecard & Starred Milestones', 'PASSED', 'Monthly productivity recognition, achievement badges, and milestone showcase')
]

for comp, feat, stat, note in items:
    row_cells = table.add_row().cells
    row_cells[0].text = comp
    row_cells[1].text = feat
    row_cells[2].text = stat
    row_cells[3].text = note

doc.save(r'c:\VARUN VK\Antigravity\orbita\Orbita_Testing_and_Debugging_Report.docx')
print('Orbita_Testing_and_Debugging_Report.docx updated successfully!')
